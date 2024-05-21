import pandas as pd
from docx import Document


def format_number2(value):
    try:
        return f"{float(value):.2f}"
    except ValueError:
        return str(value)


def format_number3(value):
    try:
        return f"{float(value):.3f}"
    except ValueError:
        return str(value)


def format_number4(value):
    try:
        return f"{float(value):.4f}"
    except ValueError:
        return str(value)


df1 = pd.read_csv('result1.csv', header=None)
df2 = pd.read_csv('result2.csv', header=None)
df3 = pd.read_csv('result3.csv', header=None)
df4 = pd.read_csv('result4.csv', header=None)
df5 = pd.read_csv('result5.csv', header=None)
df6 = pd.read_csv('result6.csv', header=None)

doc = Document()

table1 = doc.add_table(rows=1, cols=len(df1.columns))
doc.add_paragraph("\n")
table2 = doc.add_table(rows=1, cols=(len(df2.columns)+1))
doc.add_paragraph("\n")
table3 = doc.add_table(rows=1, cols=len(df3.columns)+1)
doc.add_paragraph("\n")
table4 = doc.add_table(rows=1, cols=len(df4.columns)+1)
doc.add_paragraph("\n")
table5 = doc.add_table(rows=1, cols=len(df5.columns)+1)
doc.add_paragraph("\n")
table6 = doc.add_table(rows=1, cols=len(df6.columns)+1)

names = ['Алюминий', 'Латунь', 'Сталь', 'Дерево', 'Плексиглас', 'Свинец']
column_names1 = ['№', 'Алюминий', 'Латунь', 'Сталь', 'Дерево', 'Плексиглас', 'Свинец']
column_names2 = ['Вещество', 'Плотность', 'Диаметр', 'Масса']
column_names3 = ['Вещество', 't, мс']
column_names4 = ['Вещество', 'delta t, мс']
column_names5 = ['Вещество', 'g, м/с^2']
column_names6 = ['Вещество', 'delta g, м/с^2']

header_cells1 = table1.rows[0].cells
for i, column_name in enumerate(column_names1):
    header_cells1[i].text = column_name

for index, row in df1.iterrows():
    row_cells = table1.add_row().cells
    for i, cell_value in enumerate(row):
        row_cells[i].text = format_number3(cell_value)

header_cells2 = table2.rows[0].cells
for i, column_name in enumerate(column_names2):
    header_cells2[i].text = column_name

for index, row in df2.iterrows():
    row_cells = table2.add_row().cells
    for i in range(len(row) + 1):
        if i == 0:
            row_cells[i].text = names[index]
        elif 0 < i < len(row):
            row_cells[i].text = format_number2(row[i - 1])
        else:
            row_cells[i].text = format_number4(row[i - 1])

header_cells3 = table3.rows[0].cells
for i, column_name in enumerate(column_names3):
    header_cells3[i].text = column_name

for index, row in df3.iterrows():
    row_cells = table3.add_row().cells
    for i in range (len(row) + 1):
        if i == 0:
            row_cells[i].text = names[index]
        else:
            row_cells[i].text = format_number4(row[i - 1])

header_cells4 = table4.rows[0].cells
for i, column_name in enumerate(column_names4):
    header_cells4[i].text = column_name

for index, row in df4.iterrows():
    row_cells = table4.add_row().cells
    for i in range(len(row) + 1):
        if i == 0:
            row_cells[i].text = names[index]
        else:
            row_cells[i].text = format_number4(row[i - 1])

header_cells5 = table5.rows[0].cells
for i, column_name in enumerate(column_names5):
    header_cells5[i].text = column_name

for index, row in df5.iterrows():
    row_cells = table5.add_row().cells
    for i in range(len(row) + 1):
        if i == 0:
            row_cells[i].text = names[index]
        else:
            row_cells[i].text = format_number4(row[i - 1])

header_cells6 = table6.rows[0].cells
for i, column_name in enumerate(column_names6):
    header_cells6[i].text = column_name

for index, row in df6.iterrows():
    row_cells = table6.add_row().cells
    for i in range(len(row) + 1):
        if i == 0:
            row_cells[i].text = names[index]
        else:
            row_cells[i].text = format_number4(row[i - 1])


doc_file = 'output.docx'
doc.save(doc_file)

