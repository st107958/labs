import pandas as pd
from docx import Document


def format_number4(value):
    try:
        return f"{float(value):.4f}"
    except ValueError:
        return str(value)

def format_number8(value):
    try:
        return f"{float(value):.10f}"
    except ValueError:
        return str(value)


df1 = pd.read_csv('result1.csv', header=None)
df2 = pd.read_csv('result2.csv', header=None)
df3 = pd.read_csv('result3.csv', header=None)

doc = Document()

table1 = doc.add_table(rows=1, cols=len(df1.columns))
doc.add_paragraph("\n")
table2 = doc.add_table(rows=1, cols=len(df2.columns))
doc.add_paragraph("\n")
table3 = doc.add_table(rows=1, cols=len(df3.columns))
column_names1 = ['№', 'Результаты отдельных наблюдений (U), В', 'Погрешность прибора на данной шкале (delta U), В']
column_names2 = ['№', 'Результаты отдельных наблюдений (U), В', 'Случайные отклонения от среднего - d, В', 'd^2, (В)^2']
column_names3 = ['№', 'Интервал', 'Число случаев', 'Доля']

header_cells1 = table1.rows[0].cells
for i, column_name in enumerate(column_names1):
    header_cells1[i].text = column_name

for index, row in df1.iterrows():
    row_cells = table1.add_row().cells
    for i, cell_value in enumerate(row):
        row_cells[i].text = format_number4(cell_value)

header_cells2 = table2.rows[0].cells
for i, column_name in enumerate(column_names2):
    header_cells2[i].text = column_name

for index, row in df2.iterrows():
    row_cells = table2.add_row().cells
    for i, cell_value in enumerate(row):
        row_cells[i].text = format_number8(cell_value)

header_cells3 = table3.rows[0].cells
for i, column_name in enumerate(column_names3):
    header_cells3[i].text = column_name

for index, row in df3.iterrows():
    row_cells = table3.add_row().cells
    for i, cell_value in enumerate(row):
        row_cells[i].text = format_number4(cell_value)


doc_file = 'output.docx'
doc.save(doc_file)

