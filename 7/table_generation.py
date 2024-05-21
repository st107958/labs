import pandas as pd
from docx import Document


def format_number(value):
    try:
        return f"{float(value):.5f}"
    except ValueError:
        return str(value)


df1 = pd.read_csv('result1.csv', header=None)
df2 = pd.read_csv('result2.csv', header=None)

doc = Document()

table1 = doc.add_table(rows=1, cols=len(df1.columns))
doc.add_paragraph("\n")
table2 = doc.add_table(rows=1, cols=len(df1.columns))


column_names = ['№', 'U_eb', 'U_kb', 'I_k', 'ln_I_k']

header_cells = table1.rows[0].cells
for i, column_name in enumerate(column_names):
    header_cells[i].text = column_name

for index, row in df1.iterrows():
    row_cells = table1.add_row().cells
    for i, cell_value in enumerate(row):
        row_cells[i].text = format_number(cell_value)

header_cells = table2.rows[0].cells
for i, column_name in enumerate(column_names):
    header_cells[i].text = column_name

for index, row in df2.iterrows():
    row_cells = table2.add_row().cells
    for i, cell_value in enumerate(row):
        row_cells[i].text = format_number(cell_value)

doc_file = 'output.docx'
doc.save(doc_file)

